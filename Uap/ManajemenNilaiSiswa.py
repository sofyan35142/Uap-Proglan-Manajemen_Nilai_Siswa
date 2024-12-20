from docx import Document
from docx.shared import Inches
import os

class Siswa:
    def __init__(self, nama, kelas, foto=None):
        self.nama = nama
        self.kelas = kelas
        self.nilai = {}
        self.foto = foto  # Menyimpan nama file foto siswa (bisa path)

    def tambah_nilai(self, mata_pelajaran, nilai):
        self.nilai[mata_pelajaran] = nilai

    def edit_nilai(self, mata_pelajaran, nilai):
        if mata_pelajaran in self.nilai:
            self.nilai[mata_pelajaran] = nilai
            print(f"Nilai untuk {mata_pelajaran} berhasil diperbarui menjadi {nilai}.")
        else:
            print(f"Mata pelajaran {mata_pelajaran} tidak ditemukan.")

    def lihat_nilai(self):
        return self.nilai


class ManajemenNilai:
    def __init__(self):
        self.data_siswa = []

    def tambah_siswa(self, nama, kelas, foto=None):
        siswa = Siswa(nama, kelas, foto)
        self.data_siswa.append(siswa)
        print(f"Siswa {nama} berhasil ditambahkan ke kelas {kelas}.")
        self.simpan_data_ke_docx()  # Simpan data ke DOCX setelah menambah siswa

    def edit_siswa(self):
        if not self.data_siswa:
            print("Tidak ada siswa yang terdaftar. Tambahkan siswa terlebih dahulu.")
            return

        print("\nDaftar Siswa:")
        for i, siswa in enumerate(self.data_siswa):
            print(f"{i + 1}. {siswa.nama} (Kelas {siswa.kelas})")

        try:
            pilihan = int(input("Pilih nomor siswa yang ingin diedit: "))
            if 1 <= pilihan <= len(self.data_siswa):
                siswa = self.data_siswa[pilihan - 1]
                print(f"Mengedit data untuk {siswa.nama}.")
                siswa.nama = input("Masukkan nama baru: ") or siswa.nama
                siswa.kelas = input("Masukkan kelas baru: ") or siswa.kelas
                siswa.foto = input("Masukkan foto siswa (path atau nama file): ") or siswa.foto
                print("Data siswa berhasil diperbarui.")
                self.simpan_data_ke_docx()  # Simpan data ke DOCX setelah edit siswa
            else:
                print("Nomor siswa tidak valid.")
        except ValueError:
            print("Input tidak valid. Masukkan nomor siswa dengan angka.")

    def cari_siswa(self, nama):
        for siswa in self.data_siswa:
            if siswa.nama == nama:
                return siswa
        return None

    def input_nilai(self):
        if not self.data_siswa:
            print("Tidak ada siswa yang terdaftar. Tambahkan siswa terlebih dahulu.")
            return

        print("\nDaftar Siswa:")
        for i, siswa in enumerate(self.data_siswa):
            print(f"{i + 1}. {siswa.nama} (Kelas {siswa.kelas})")

        try:
            pilihan = int(input("Pilih nomor siswa: "))
            if 1 <= pilihan <= len(self.data_siswa):
                siswa = self.data_siswa[pilihan - 1]
                print(f"\nNilai saat ini untuk {siswa.nama}:")
                for mata_pelajaran_existing, nilai_existing in siswa.lihat_nilai().items():
                    print(f"{mata_pelajaran_existing}: {nilai_existing}")

                mata_pelajaran = input("Masukkan mata pelajaran: ")
                nilai = float(input("Masukkan nilai: "))
                siswa.tambah_nilai(mata_pelajaran, nilai)
                print(f"Nilai {mata_pelajaran} berhasil ditambahkan untuk {siswa.nama}.")
                self.simpan_data_ke_docx()  # Simpan data ke DOCX setelah input nilai
            else:
                print("Nomor siswa tidak valid.")
        except ValueError:
            print("Input tidak valid. Masukkan nomor siswa dengan angka.")

    def edit_nilai(self):
        if not self.data_siswa:
            print("Tidak ada siswa yang terdaftar. Tambahkan siswa terlebih dahulu.")
            return

        print("\nDaftar Siswa:")
        for i, siswa in enumerate(self.data_siswa):
            print(f"{i + 1}. {siswa.nama} (Kelas {siswa.kelas})")

        try:
            pilihan = int(input("Pilih nomor siswa: "))
            if 1 <= pilihan <= len(self.data_siswa):
                siswa = self.data_siswa[pilihan - 1]
                print(f"\nNilai saat ini untuk {siswa.nama}:")
                for mata_pelajaran_existing, nilai_existing in siswa.lihat_nilai().items():
                    print(f"{mata_pelajaran_existing}: {nilai_existing}")

                mata_pelajaran = input("Masukkan mata pelajaran yang ingin diedit: ")
                if mata_pelajaran in siswa.nilai:
                    nilai = float(input("Masukkan nilai baru: "))
                    siswa.edit_nilai(mata_pelajaran, nilai)
                    self.simpan_data_ke_docx()  # Simpan data ke DOCX setelah edit nilai
                else:
                    print(f"Mata pelajaran {mata_pelajaran} tidak ditemukan.")
            else:
                print("Nomor siswa tidak valid.")
        except ValueError:
            print("Input tidak valid. Masukkan nomor siswa dengan angka.")

    def lihat_semua_nilai(self):
        if not self.data_siswa:
            print("Tidak ada data siswa.")
        for siswa in self.data_siswa:
            print(f"\nNama: {siswa.nama}, Kelas: {siswa.kelas}")
            if siswa.foto:
                print(f"Foto: {siswa.foto}")
            for mata_pelajaran, nilai in siswa.lihat_nilai().items():
                print(f"{mata_pelajaran}: {nilai}")

    def simpan_data_ke_docx(self, filename="data_siswa.docx"):
        doc = Document()
        doc.add_heading("Data Siswa", 0)

        # Menambahkan tabel untuk setiap siswa
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Nama'
        hdr_cells[1].text = 'Kelas'
        hdr_cells[2].text = 'Foto'
        hdr_cells[3].text = 'Mata Pelajaran'
        hdr_cells[4].text = 'Nilai'

        # Menambahkan data siswa ke tabel
        for siswa in self.data_siswa:
            row_cells = table.add_row().cells
            row_cells[0].text = siswa.nama
            row_cells[1].text = siswa.kelas
            # Mata Pelajaran hanya nama mata pelajaran saja, tanpa nilai
            row_cells[3].text = "\n".join([mata_pelajaran for mata_pelajaran in siswa.lihat_nilai().keys()]) or "Tidak ada mata pelajaran"
            # Nilai berada di kolom Nilai
            row_cells[4].text = "\n".join([str(nilai) for nilai in siswa.lihat_nilai().values()]) or "Tidak ada nilai"

            # Menambahkan foto jika ada
            if siswa.foto and os.path.exists(siswa.foto):
                try:
                    # Menambahkan foto ke dalam tabel
                    cell = row_cells[2]
                    cell.text = ""  # Clear the text of the cell
                    cell_paragraph = cell.add_paragraph()
                    cell_paragraph.add_run().add_picture(siswa.foto, width=Inches(1.0))
                except Exception as e:
                    print(f"Error menambahkan foto: {e}")
                    row_cells[2].text = "Gagal menambahkan foto"
            else:
                row_cells[2].text = "Tidak ada foto"

        doc.save(filename)
        print(f"Data berhasil disimpan ke {filename}")


# Fungsi utama untuk menjalankan aplikasi
def main():
    manajemen = ManajemenNilai()

    while True:
        print("\n=== Aplikasi Manajemen Nilai Siswa ===")
        print("1. Tambah Siswa")
        print("2. Edit Siswa")
        print("3. Input Nilai Siswa")
        print("4. Edit Nilai Siswa")
        print("5. Lihat Semua Nilai Siswa")
        print("6. Keluar")

        pilihan = input("Pilih menu (1/2/3/4/5/6): ")

        if pilihan == "1":
            nama = input("Masukkan nama siswa: ")
            kelas = input("Masukkan kelas siswa: ")
            foto = input("Masukkan foto siswa (path atau nama file): ")
            manajemen.tambah_siswa(nama, kelas, foto)
        elif pilihan == "2":
            manajemen.edit_siswa()
        elif pilihan == "3":
            manajemen.input_nilai()
        elif pilihan == "4":
            manajemen.edit_nilai()
        elif pilihan == "5":
            manajemen.lihat_semua_nilai()
        elif pilihan == "6":
            print("Terima kasih telah menggunakan aplikasi ini.")
            break
        else:
            print("Pilihan tidak valid, silakan coba lagi.")


if __name__ == "__main__":
    main()
